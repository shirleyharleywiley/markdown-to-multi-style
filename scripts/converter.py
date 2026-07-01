#!/usr/bin/env python3
"""Markdown 样式转换器 — 8 种版式，输出 Word 或 HTML"""

import os, sys, re, json, markdown
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from bs4 import BeautifulSoup

_SELF_DIR = os.path.dirname(os.path.abspath(__file__))
_DATA = None


def _load():
    global _DATA
    if _DATA is None:
        with open(os.path.join(_SELF_DIR, 'data.json'), encoding='utf-8') as f:
            raw = json.load(f)
        for pk, pv in raw['palettes'].items():
            for k in ('primary', 'primary_dark', 'primary_light', 'secondary', 'secondary_dark', 'secondary_light'):
                if k in pv:
                    pv[k] = tuple(pv[k])
        _DATA = raw
    return _DATA


PALETTES = property(lambda self: _load()['palettes'])
STYLES   = property(lambda self: _load()['styles'])


class _PALETTES(dict):
    def __getitem__(self, k): return _load()['palettes'][k]
    def keys(self): return _load()['palettes'].keys()
    def __contains__(self, k): return k in _load()['palettes']

class _STYLES(dict):
    def __getitem__(self, k): return _load()['styles'][k]
    def keys(self): return _load()['styles'].keys()
    def __contains__(self, k): return k in _load()['styles']

PALETTES = _PALETTES()
STYLES   = _STYLES()


def build_color_dict(style_key):
    """根据版式（使用默认配色），构建 docx_colors 字典"""
    data = _load()
    style = data['styles'][style_key]
    palette_key = style.get('default_palette', 'blue')
    pal = data['palettes'][palette_key]
    cmap = style["color_map"]

    lookup = {
        "PRIMARY": pal["primary"],
        "PRIMARY_DARK": pal["primary_dark"],
        "PRIMARY_LIGHT": pal["primary_light"],
        "SECONDARY": pal["secondary"],
        "SECONDARY_DARK": pal["secondary_dark"],
        "SECONDARY_LIGHT": pal["secondary_light"],
        "TEXT": (0x33, 0x33, 0x33),
        "TEXT_DARK": (0x1a, 0x1a, 0x1a),
        "TEXT_WHITE": (0xff, 0xff, 0xff),
        "CODE_BG": (0xf5, 0xf5, 0xf5),
        "CODE_TEXT": (0xe9, 0x1e, 0x63),
        "TABLE_BORDER": (0xe0, 0xe0, 0xe0),
        "TABLE_ALT": (0xf8, 0xf8, 0xf8),
        "TABLE_HEADER_BG": (0xff, 0xff, 0xff),
        "BLOCKQUOTE_TEXT": (0x55, 0x55, 0x55),
        "BLOCKQUOTE_BORDER": (0xdd, 0xdd, 0xdd),
    }

    result = {}
    for role, key in cmap.items():
        result[role] = lookup.get(key, (0x33, 0x33, 0x33))
    return result


def css_with_palette(style_key):
    """将版式 CSS 中的占位符替换为实际配色值（使用默认配色）"""
    data = _load()
    style = data['styles'][style_key]
    palette_key = style.get('default_palette', 'blue')
    pal = data['palettes'][palette_key]

    def c(key):
        p = pal[key]
        return f"#{p[0]:02x}{p[1]:02x}{p[2]:02x}"

    adj_key = pal["adjacent"]
    adj_pal = data['palettes'][adj_key]
    adj_hex = f"#{adj_pal['primary'][0]:02x}{adj_pal['primary'][1]:02x}{adj_pal['primary'][2]:02x}"

    replacements = {
        "{PRIMARY}": c("primary"),
        "{PRIMARY_DARK}": c("primary_dark"),
        "{PRIMARY_LIGHT}": c("primary_light"),
        "{SECONDARY}": c("secondary"),
        "{SECONDARY_DARK}": c("secondary_dark"),
        "{SECONDARY_LIGHT}": c("secondary_light"),
        "{ADJACENT}": adj_hex,
        "{TERTIARY}": c("primary_dark"),
        "{TABLE_BORDER}": "#e0e0e0",
        "{TABLE_ALT}": "#f8f8f8",
        "{TABLE_HEADER_BG}": "#ffffff",
        "{CODE_BG}": "#f5f5f5",
        "{CODE_TEXT}": "#e91e63",
        "{BLOCKQUOTE_TEXT}": "#555555",
        "{BLOCKQUOTE_BORDER}": "#dddddd",
    }

    css = style["html_css"]
    for placeholder, value in replacements.items():
        css = css.replace(placeholder, value)
    return css


# ============================================================
# 核心工具函数
# ============================================================

def is_separator_row(line):
    """判断是否是表格分隔行"""
    stripped = line.strip()
    if not stripped.startswith('|'):
        return False
    inner = stripped.strip('|')
    cells = [c.strip() for c in inner.split('|')]
    return all(re.match(r'^[-: ]+$', c) for c in cells)

def is_valid_table_row(line):
    """判断是否是有效的表格数据行（非分隔行，有多列内容）"""
    stripped = line.strip()
    if not stripped.startswith('|') or is_separator_row(stripped):
        return False
    inner = stripped.strip('|')
    cells = [c.strip() for c in inner.split('|')]
    return len(cells) >= 2 and any(c for c in cells)

def has_chinese(text):
    """检查文本是否包含中文"""
    return any('一' <= c <= '鿿' for c in text)

def has_english(text):
    """检查文本是否包含英文"""
    return any('a' <= c <= 'z' or 'A' <= c <= 'Z' for c in text)

def is_bilingual_merge(row1, row2):
    """判断两行是否应该合并（双语对照场景）"""
    cells1 = [c.strip() for c in row1.strip().strip('|').split('|')]
    cells2 = [c.strip() for c in row2.strip().strip('|').split('|')]

    if len(cells1) != len(cells2) or len(cells1) < 2:
        return False

    for c1, c2 in zip(cells1, cells2):
        if not c1 or not c2:
            continue
        has_ch1, has_ch2 = has_chinese(c1), has_chinese(c2)
        has_en1, has_en2 = has_english(c1), has_english(c2)
        is_chinese_english_pair = (has_ch1 and not has_ch2 and has_en2 and not has_en1) or \
                                  (has_ch2 and not has_ch1 and has_en1 and not has_en2)
        if not is_chinese_english_pair:
            return False

    return True

def fix_bilingual_tables(md_text):
    """双语对照表格合并 — 仅对真正的中英对照行进行合并，普通表格不受影响"""
    lines = md_text.split('\n')
    result = []
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if is_separator_row(stripped):
            result.append(line)
            i += 1
            continue

        if is_valid_table_row(stripped):
            if i + 2 < len(lines):
                sep_line = lines[i + 1]
                next_content = lines[i + 2]

                if is_separator_row(sep_line) and is_valid_table_row(next_content):
                    if is_bilingual_merge(stripped, next_content):
                        cells1 = [c.strip() for c in stripped.strip('|').split('|')]
                        cells2 = [c.strip() for c in next_content.strip('|').split('|')]
                        merged = '|'.join([(c1 + ' / ' + c2) if c1 != c2 else c1
                                          for c1, c2 in zip(cells1, cells2)])
                        result.append('| ' + merged + ' |')
                        i += 3
                        continue

            result.append(line)
        else:
            result.append(line)

        i += 1

    return '\n'.join(result)


def fix_table_preceding_blank_lines(md_text):
    """修复表格前缺少空行的问题 — 仅在表格开始处插入空行（不包括表格内容之间）"""
    lines = md_text.split('\n')
    result = []
    in_table = False  # 是否在表格内

    for i, line in enumerate(lines):
        stripped = line.strip()

        # 检测到表格行
        if stripped.startswith('|'):
            inner = stripped.strip('|')
            cells = [c.strip() for c in inner.split('|')]

            # 是有效表格行（非分隔行，有多列内容）
            if cells and len(cells) >= 2 and not all(re.match(r'^[-: ]+$', c) for c in cells):
                # 表格的第一个有效行前需要空行
                if not in_table:
                    # 检查前一行是否为空行
                    if result and result[-1].strip() != '':
                        result.append('')
                    in_table = True
            else:
                # 是分隔行，继续表格
                pass

            result.append(line)
        else:
            # 非表格行，退出表格状态
            in_table = False
            result.append(line)

    return '\n'.join(result)


def fix_special_syntax(md_text):
    """引用块内列表语法处理（暂不使用）"""
    return md_text


def _cjk_width(text):
    """估算单行文本渲染宽度（em）：
    - CJK 字符按 1em 计
    - 拉丁/数字字符按 0.5em 计
    返回 em 值
    """
    cjk = sum(1 for c in text if '一' <= c <= '鿿')
    ascii_count = len(text) - cjk
    return cjk * 1.0 + ascii_count * 0.5


def apply_table_layout(html):
    """统一表格列宽规则（按列名匹配 + 总宽度约束 + 总高度最矮）：

    核心设计原则：
    - 列宽按"该列内容字数"的**平方根**比例分配（而非线性比例）：
        w[i] ∝ sqrt(n[i] / n_max)
      字数最多的列宽度 / 字数次多列宽度的平方 ≈ 二者字数之比。
      这样能避免"字数最多的列无限制吞噬宽度"，让其他列也有合理宽度。
    - 必须满足"总宽度 ≤ 容器宽度"（CONTAINER_EM=60em）。
    - 表格内联样式：table-layout:fixed + word-break:break-word + white-space:normal
      （让单元格内容在列宽内自动换行 → 高度最小化）

    规则：
    1. 同表头同列同宽（按列名匹配）
    2. 列宽按该列内容总字数（em 维度）的平方根比例分配
    3. 总宽 ≤ 60em，超出则等比缩放
    4. 每列最低保证 6em

    实现方式：
    - 两遍扫描：
      · 第一遍：收集"列名 → (表头 em, 该列内容总 em)" 用于后续 sqrt 缩放
      · 第二遍：按列名回填每张表的 colgroup
    """
    import math
    # 容器宽度（em），约等于 .md-body 的 max-width:860px / font-size:14px
    CONTAINER_EM = 60
    MIN_COL_EM = 6  # 每列最低保证宽度
    BASE_EM = 8     # 最小列宽基准（与原 _columns_em 的 base 一致）

    soup = BeautifulSoup(html, 'html.parser')
    tables = soup.find_all('table')
    if not tables:
        return html

    def _get_ths(table):
        thead = table.find('thead')
        if thead is not None:
            ths = thead.find_all('th')
            if ths:
                return ths
        first_tr = table.find('tr')
        return first_tr.find_all('th') if first_tr else []

    def _get_tds_by_col(table, n_cols):
        """返回 [[col0_tds_text...], [col1_tds_text...], ...]，仅取 tbody 行

        采样策略：每列最多取前 SAMPLE_LIMIT 个单元（不需要全表精确）。
        """
        SAMPLE_LIMIT = 5  # 每列最多采样 5 个单元，足够反映宽度需求

        result = [[] for _ in range(n_cols)]
        tbody = table.find('tbody')
        rows = tbody.find_all('tr') if tbody else []
        if not rows:
            all_trs = table.find_all('tr')
            rows = all_trs[1:] if all_trs else []
        for tr in rows:
            cells = tr.find_all(['td', 'th'])
            for idx, cell in enumerate(cells[:n_cols]):
                if len(result[idx]) < SAMPLE_LIMIT:
                    result[idx].append(cell.get_text(strip=True))
                else:
                    # 该列已采满，跳过（节省时间）
                    continue
        return result

    # 第一遍：收集"列名 → (表头 em, 该列内容总 em, 是否短列, 该列最大单格 em)"
    # 跨表合并策略：把所有表的样本合在一起算总 em / max_single / all_short
    SHORT_THRESHOLD = 15  # 单格内容 cjk 宽度低于此值视为"短列"
    SHORT_FIXED_EM = 6   # 短列固定宽度

    # 用嵌套 dict 收集所有样本：col_samples[name] = [all texts from all tables]
    col_samples = {}

    col_data = {}
    for table in tables:
        ths = _get_ths(table)
        if not ths:
            continue
        col_names = [th.get_text(strip=True) for th in ths]
        tds_by_col = _get_tds_by_col(table, len(col_names))
        for idx, name in enumerate(col_names):
            if not name:
                continue
            header_em = _cjk_width(name)
            samples = tds_by_col[idx] if idx < len(tds_by_col) else []
            # 累加所有样本到池中
            col_samples.setdefault(name, []).extend(samples)
            if name not in col_data:
                col_data[name] = (header_em, 0, True, 0)  # 临时值
            else:
                old_header, _, _, old_max = col_data[name]
                col_data[name] = (max(old_header, header_em), 0, True, old_max)

    # 第二遍合并：基于完整样本池计算最终值
    for name, samples in col_samples.items():
        header_em = col_data[name][0]
        body_total_em = sum(_cjk_width(s) for s in samples)
        max_single_em = max((_cjk_width(s) for s in samples), default=0)
        all_short = bool(samples) and all(_cjk_width(s) < SHORT_THRESHOLD for s in samples)
        col_data[name] = (header_em, body_total_em, all_short, max_single_em)

    # 第二遍：按"字数平方根"比例分配长列预算
    # 核心算法：
    # - 短列固定 6em，不参与预算分配
    # - 长列预算 = CONTAINER_EM - 短列总宽
    # - 各长列宽度 = 预算 × sqrt(n_i / n_max) / Σsqrt(n_i / n_max)
    #   这样 W_最多 / W_次多 的平方 = n_最多 / n_次多（你定的核心规则）
    #   例：n_max=120, n_2nd=30 → W_max / W_2nd = √(120/30) = 2
    col_name_width = {}
    if col_data:
        # 短列直接固定 6em
        long_cols = {n: v for n, v in col_data.items() if not v[2]}
        for name in col_data:
            if col_data[name][2]:
                col_name_width[name] = SHORT_FIXED_EM

        # 长列按 sqrt 字数比分配预算
        if long_cols:
            long_total_chars = {n: v[1] for n, v in long_cols.items()}
            max_total_chars = max(long_total_chars.values()) if long_total_chars else 0

            if max_total_chars > 0:
                # 计算 sqrt 权重
                sqrt_weights = {n: math.sqrt(c / max_total_chars)
                                for n, c in long_total_chars.items()}
                total_sqrt = sum(sqrt_weights.values())

                # 长列总预算
                short_total = sum(SHORT_FIXED_EM for n in col_data if col_data[n][2])
                long_budget = max(0, CONTAINER_EM - short_total)

                if total_sqrt > 0 and long_budget > 0:
                    for name in long_cols:
                        col_name_width[name] = max(
                            MIN_COL_EM,
                            long_budget * sqrt_weights[name] / total_sqrt
                        )
                else:
                    for name in long_cols:
                        col_name_width[name] = MIN_COL_EM
            else:
                for name in long_cols:
                    col_name_width[name] = MIN_COL_EM

        # 兜底：总宽仍超 CONTAINER_EM（MIN_COL_EM 兜底后）→ 等比再缩
        total_em = sum(col_name_width.values())
        if total_em > CONTAINER_EM:
            scale = CONTAINER_EM / total_em
            for name in col_name_width:
                col_name_width[name] *= scale

    # 第三遍：按列名回填每张表的 colgroup
    for table in tables:
        old = table.find('colgroup')
        if old:
            old.decompose()

        ths = _get_ths(table)
        if not ths:
            continue

        colgroup = soup.new_tag('colgroup')
        for th in ths:
            name = th.get_text(strip=True)
            col = soup.new_tag('col')
            if name in col_name_width:
                em_val = round(col_name_width[name], 1)
                col['style'] = f'width:{em_val}em'
            else:
                col['style'] = 'width:auto'
            colgroup.append(col)
        table.insert(0, colgroup)

        existing = table.get('style', '')
        if 'table-layout' not in existing:
            table['style'] = (
                existing + '; table-layout:fixed; word-break:break-word; white-space:normal'
            ).strip('; ')

    return str(soup)


def md_to_html(md_path):
    md = open(md_path, encoding='utf-8').read()
    fixed = fix_bilingual_tables(md)
    fixed = fix_table_preceding_blank_lines(fixed)
    fixed = fix_special_syntax(fixed)
    html = markdown.Markdown(extensions=['fenced_code', 'tables', 'nl2br', 'sane_lists', 'toc']).convert(fixed)
    html = apply_table_layout(html)
    return html


def md_to_docx(md_path):
    md = open(md_path, encoding='utf-8').read()
    fixed = fix_bilingual_tables(md)
    fixed = fix_table_preceding_blank_lines(fixed)
    fixed = fix_special_syntax(fixed)
    html = markdown.Markdown(extensions=['fenced_code', 'tables', 'nl2br', 'sane_lists']).convert(fixed)
    soup = BeautifulSoup(html, 'html.parser')
    return soup


# ============================================================
# HTML 生成
# ============================================================

def generate_html(md_path, style_key, output_path=None):
    html_content = md_to_html(md_path)
    css = css_with_palette(style_key)

    html = (
        '<!DOCTYPE html>\n'
        '<html lang="zh-CN">\n'
        '<head>\n'
        '<meta charset="UTF-8">\n'
        '<meta name="viewport" content="width=device-width, initial-scale=1.0">\n'
        '<title>' + os.path.basename(md_path) + '</title>\n'
        '<style>\n'
        '  * { box-sizing: border-box; margin: 0; padding: 0; }\n'
        '  ' + css + '\n'
        '  .md-body img { max-width: 100%; border-radius: 5px; display: block; margin: 16px 0; }\n'
        '  .md-body pre::-webkit-scrollbar { height: 6px; }\n'
        '  .md-body pre::-webkit-scrollbar-track { background: #1a1a1a; }\n'
        '  .md-body pre::-webkit-scrollbar-thumb { background: #444; border-radius: 3px; }\n'
        '</style>\n'
        '</head>\n'
        '<body>\n'
        '<div class="md-body">\n'
        + html_content + '\n'
        '</div>\n'
        '</body>\n'
        '</html>'
    )

    if output_path is None:
        base = os.path.splitext(md_path)[0]
        output_path = base + '-' + style_key + '.html'
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(html)
    return output_path


# ============================================================
# Word 生成
# ============================================================

def set_eastasia_font(run, name='PingFang SC'):
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), name)
    for existing in rPr.findall(qn('w:rFonts')):
        rPr.remove(existing)
    rPr.insert(1, rFonts)


def make_run(p, text, bold=False, italic=False, font_name='PingFang SC',
             font_size=11, color=None):
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    set_eastasia_font(run, font_name)
    return run


def set_cell_borders(cell, color='ddd'):
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement('w:' + side)
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _fmt(c):
    if isinstance(c, str):
        return c.upper()
    return f"{c[0]:02X}{c[1]:02X}{c[2]:02X}"


def add_table_docx(doc, table, colors):
    headers = []
    thead = table.find('thead')
    if thead:
        for th in thead.find_all('th'):
            headers.append(th.get_text(strip=True))
    if not headers:
        first_tr = table.find('tr')
        if first_tr:
            headers = [th.get_text(strip=True) for th in first_tr.find_all('th')]

    rows_data = []
    tbody = table.find('tbody')
    if tbody:
        trs = tbody.find_all('tr')
    else:
        all_trs = table.find_all('tr')
        trs = all_trs[1:] if headers and len(all_trs) > 1 else all_trs
    for tr in trs:
        cells = tr.find_all(['th', 'td'])
        row = [c.get_text(strip=True) for c in cells]
        if row and any(c for c in row):
            rows_data.append(row)

    if not headers or not rows_data:
        return

    hdr_bg = _fmt(colors["table_header_bg"])
    alt_bg = _fmt(colors["table_alt_bg"])
    border = _fmt(colors["table_border"])
    hdr_text = (0xff, 0xff, 0xff)
    cell_text = colors["text"]

    t = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    t.style = 'Table Grid'

    hdr = t.rows[0]
    for j, h in enumerate(headers):
        cell = hdr.cells[j]
        p = cell.paragraphs[0]; p.clear()
        make_run(p, h, bold=True, font_size=10.5, color=hdr_text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tcPr = cell._element.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hdr_bg)
        tcPr.append(shd)
        set_cell_borders(cell, border)

    for i, row in enumerate(rows_data):
        tr = t.rows[i + 1]
        for j, cell_text2 in enumerate(row):
            if j >= len(headers): continue
            cell = tr.cells[j]
            p = cell.paragraphs[0]; p.clear()
            make_run(p, cell_text2, font_size=10, color=cell_text)
            tcPr = cell._element.get_or_add_tcPr()
            fill = 'FFFFFF' if i % 2 == 0 else alt_bg
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            tcPr.append(shd)
            set_cell_borders(cell, border)

    doc.add_paragraph()


def add_hr_docx(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'E8E8E8')
    pBdr.append(bottom); pPr.append(pBdr)


def generate_docx(md_path, style_key, output_path=None):
    soup = md_to_docx(md_path)
    colors = build_color_dict(style_key)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    def add_h(tag):
        level = int(tag.name[1])
        text = tag.get_text(strip=True)
        p = doc.add_heading('', level=level)
        p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
        p.paragraph_format.space_after = Pt(8)
        c = colors["h1"] if level == 1 else colors["h2"] if level == 2 else colors["text"]
        make_run(p, text, bold=True,
                 font_size=20 if level == 1 else 14 if level == 2 else 13,
                 color=c)
        if level <= 2:
            pPr = p._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            b = OxmlElement('w:bottom')
            b.set(qn('w:val'), 'none'); b.set(qn('w:sz'), '0')
            b.set(qn('w:space'), '0'); b.set(qn('w:color'), 'auto')
            pBdr.append(b); pPr.append(pBdr)

    def add_p(tag):
        text = tag.get_text(strip=True)
        if not text: return
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        make_run(p, text, font_size=11, color=colors["text"])

    def add_blockquote(tag):
        lines = []
        for child in tag.children:
            if child.name == 'br': lines.append('\n')
            else: lines.append(child.get_text())
        text = ''.join(lines).strip()
        for line in text.split('\n'):
            if line.strip():
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.left_indent = Cm(0.5)
                pPr = p._element.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                left = OxmlElement('w:left')
                left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '12')
                left.set(qn('w:space'), '4'); left.set(qn('w:color'), _fmt(colors["blockquote_border"]))
                pBdr.append(left); pPr.append(pBdr)
                make_run(p, line.strip(), font_size=10, color=colors["blockquote_text"])

    def add_code_block(tag):
        text = tag.get_text()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        pPr = p._element.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), _fmt(colors["code_bg"]))
        pPr.append(shd)
        make_run(p, text, font_name='Courier', font_size=9.5, color=colors["code_text"])

    def add_list(tag, ordered=False):
        for item in tag.find_all('li', recursive=False):
            text = item.get_text(strip=True)
            if not text: continue
            p = doc.add_paragraph(style='List Bullet' if not ordered else 'List Number')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(0.6)
            make_run(p, text, font_size=11, color=colors["text"])

    body = soup.find('body') or soup
    for child in body.children:
        if child.name is None: continue
        tag_name = child.name.lower()
        if tag_name.startswith('h') and len(tag_name) == 2 and tag_name[1].isdigit():
            add_h(child)
        elif tag_name == 'hr':
            add_hr_docx(doc)
        elif tag_name == 'table':
            add_table_docx(doc, child, colors)
        elif tag_name == 'blockquote':
            add_blockquote(child)
        elif tag_name == 'pre':
            add_code_block(child)
        elif tag_name == 'ul':
            add_list(child, ordered=False)
        elif tag_name == 'ol':
            add_list(child, ordered=True)
        elif tag_name == 'p':
            add_p(child)

    if output_path is None:
        base = os.path.splitext(md_path)[0]
        output_path = base + '-' + style_key + '.docx'
    doc.save(output_path)
    return output_path


# ============================================================
# 主入口
# ============================================================

if __name__ == '__main__':
    md_path = sys.argv[1] if len(sys.argv) > 1 else 'README.md'
    fmt = sys.argv[2] if len(sys.argv) > 2 else 'html'
    style_key = sys.argv[3] if len(sys.argv) > 3 else 'block'

    if fmt == 'html':
        out = generate_html(md_path, style_key)
        print('HTML:', out)
    else:
        out = generate_docx(md_path, style_key)
        print('Word:', out)